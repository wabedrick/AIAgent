
import os
import platform
from docx import Document
from dotenv import load_dotenv
from duckduckgo_search import DDGS
from tavily import TavilyClient # type: ignore
import re

# Import native Google ADK requirements
from google.adk.agents import Agent
from google.adk.models.google_llm import Gemini
from google.adk.runners import InMemoryRunner
from google.adk.tools import FunctionTool, AgentTool
# from google.genai import types
from google.adk.models.lite_llm import LiteLlm
import litellm


litellm._turn_on_debug()  # type: ignore # shows full error details in Render logs

load_dotenv()

# ─── Model Configuration ─────────────────────────────────────────────────────

GROQ_MODEL = "groq/llama-3.3-70b-versatile"
GEMINI_MODEL = "gemini-2.0-flash"  # highest free tier: 1500 req/day

# Track which model is currently active
current_model_index = 0

models = [
    {
        "name": "Groq",
        "model": LiteLlm(
            model=GROQ_MODEL,
            api_key=os.getenv("GROQ_API_KEY"),
            num_retries=2,
            timeout=120
        )
    },
    {
        "name": "Gemini",
        "model": Gemini(model=GEMINI_MODEL)
    }
]

# groq_model = LiteLlm(
#     model="groq/llama-3.3-70b-versatile",
#     api_key=os.getenv("GROQ_API_KEY"),
#     num_retries=2,
#     timeout=120
# )

# gemini_model = Gemini(model="gemini-2.0-flash")

# # Start with Groq as primary
# my_model = groq_model

def get_active_model():
    """Returns the currently active model."""
    return models[current_model_index]["model"]

def switch_model():
    """Switch to the next available model."""
    global current_model_index
    previous = models[current_model_index]["name"]
    current_model_index = (current_model_index + 1) % len(models)
    active = models[current_model_index]["name"]
    print(f"[Model Router] Switched from {previous} → {active}", flush=True)

def get_model_with_fallback():
    """
    Returns a model-aware wrapper. If the active model hits a rate limit,
    it automatically switches to the other one.
    """
    return SmartModel()

# ─── Smart Model Wrapper ──────────────────────────────────────────────────────

class SmartModel:
    """
    Wraps LiteLlm/Gemini and intercepts rate limit errors to switch models.
    ADK agents accept this as a valid model since we implement __call__.
    """
    def __getattr__(self, name):
        return getattr(get_active_model(), name)

my_model = get_active_model()

# my_model = LiteLlm(
#     model="groq/llama-3.3-70b-versatile",
#     api_key=os.getenv("GROQ_API_KEY"),
#     num_retries=5,
#     timeout=60
# )

# 2. Initialize the native Google Gemini model configuration.
# The underlying Google GenAI SDK automatically grabs your API key from the environment.
# my_model = Gemini(model=my_model_name)

# Keywords that signal the user wants current/recent information
CURRENT_EVENT_KEYWORDS = [
    "today", "latest", "current", "recent", "now", "breaking",
    "this week", "this month", "this year", "2024", "2025",
    "news", "update", "happening", "live", "right now", "just",
    "yesterday", "tonight", "this morning", "new", "announced",
    "released", "launched", "trending"
]

def is_current_event_query(query: str) -> bool:
    """Detect if the query is about current or recent events."""
    query_lower = query.lower()
    return any(keyword in query_lower for keyword in CURRENT_EVENT_KEYWORDS)

def search_with_duckduckgo(query: str) -> str:
    """Fallback search using DuckDuckGo for general/non-current queries."""
    try:
        ddgs = DDGS()
        results = ddgs.text(query, max_results=3)
        if not results:
            return "No results found."
        return "\n\n".join([
            f"Title: {res['title']}\nSnippet: {res['body']}"
            for res in results
        ])
    except Exception as e:
        return f"DuckDuckGo search failed: {str(e)}"

def search_with_tavily(query: str) -> str:
    """Search using Tavily for current events and real-time data."""
    try:
        client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
        response = client.search(
            query=query,
            search_depth="advanced",
            max_results=3,
            include_answer=True,
        )
        output = ""
        if response.get("answer"):
            output += f"Direct Answer: {response['answer']}\n\n"
        results = response.get("results", [])
        if results:
            output += "\n\n".join([
                f"Title: {r['title']}\nURL: {r['url']}\nSnippet: {r['content']}"
                for r in results
            ])
        return output if output else "No results found."
    except Exception as e:
        error_msg = str(e).lower()
        # Tavily quota exceeded — fall back to DuckDuckGo
        if any(term in error_msg for term in ["quota", "limit", "exceeded", "402", "429"]):
            print("Tavily quota exceeded, falling back to DuckDuckGo.", flush=True)
            return search_with_duckduckgo(query)
        return f"Tavily search failed: {str(e)}"

def web_search(query: str) -> str:
    """
    Smart search router:
    - Current/recent event queries → Tavily (real-time, accurate)
    - General/historical queries → DuckDuckGo
    - Tavily quota exceeded → automatically falls back to DuckDuckGo
    """
    if is_current_event_query(query):
        print(f"[Router] Current event detected → using Tavily for: {query}", flush=True)
        return search_with_tavily(query)
    else:
        print(f"[Router] General query detected → using DuckDuckGo for: {query}", flush=True)
        return search_with_duckduckgo(query)

my_search_tool = FunctionTool(web_search)


# def web_search(query: str) -> str:
#     """Use this tool to search the internet for current events, facts, or data."""
#     try:
#         # Initialize DuckDuckGo search
#         ddgs = DDGS()
#         # Get the top 3 results
#         results = ddgs.text(query, max_results=3)
        
#         # Format the results into a string for the LLM to read
#         formatted_results = "\n\n".join(
#             [f"Title: {res['title']}\nSnippet: {res['body']}" for res in results]
#         )
#         return formatted_results if formatted_results else "No results found."
#     except Exception as e:
#         return f"Search failed: {str(e)}"

# # Wrap the standard function into an ADK-compatible tool
# my_search_tool = FunctionTool(web_search)

# Research Agent: Its job is to use the custom web_search tool and present findings.
research_agent = Agent(
    name="ResearchAgent",
    model=my_model,
    instruction="""You are a specialized research agent. Your only job is to use the
    web_search tool to find 2-3 pieces of relevant information on the given topic and present the findings with citations.""",
    tools=[my_search_tool], 
    output_key="research_findings",  # Stored in the session state
)

# Summarizer Agent: Its job is to summarize the text it receives.
summarizer_agent = Agent(
    name="SummarizerAgent",
    model=my_model,
    instruction="""Read the provided research findings: {research_findings}
Create a concise summary as a bulleted list with 3-5 key points.""",
    output_key="final_summary",
)

# This function will be used by the CVStylistAgent to save the generated markdown as a Word document.
def save_to_word(text_content: str, filename: str = "Document.docx") -> str:
    """
    Creates a formatted Word document from markdown text.
    
    CRITICAL INSTRUCTION FOR LLM:
    - `text_content`: The FULL markdown text.
    - `filename`: Name the file as '{FirstName}_{LastName}_{DocumentType}.docx'
      where DocumentType is one of: CV, Resume, Cover_Letter
      e.g. 'John_Doe_Cover_Letter.docx', 'John_Doe_Resume.docx'
    """
    doc = Document()
    
    lines = text_content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped.replace("## ", ""), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped.replace("### ", ""), level=2)
        elif stripped.startswith("* ") or stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    base, ext = os.path.splitext(filename)
    counter = 1
    final_filename = filename
    
    while True:
        try:
            doc.save(final_filename)
            break
        except PermissionError:
            final_filename = f"{base}_{counter}{ext}"
            counter += 1

    try:
        if platform.system() == "Windows":
            os.startfile(final_filename)
    except Exception:
        pass

    return f"Success! Document saved as '{final_filename}'."

# Wrap it for the agent
my_doc_tool = FunctionTool(save_to_word)

# CV Stylist Agent: Its job is to take the research findings, draft a CV in markdown, and 
# save it using the tool.
cv_stylist_agent = Agent(
    name="DocumentStylistAgent",
    model=my_model,
    instruction="""You are a professional career document formatter. You can create CVs, Resumes, and Cover Letters.

1. Identify the document type requested: CV, Resume, or Cover Letter.
2. Try to extract the person's full name from the research data.
   - If a clear full name is found: use Firstname_Lastname as the filename prefix.
   - If no name is found or it is unclear: use the filename prefix My instead e.g. My_CV.docx
3. Search the web using my_search_tool for the best and most upto date format for that particular document such that you give the user 
the most professional upto date and official document.
   -If it's East African format or standard, search it on the web and give according to it
   -If it's USA or Europe standards, you can get all these on the internet. Please use your knowledge base on the format and standard, such it
   on the web.
4. Draft the complete, highly detailed Markdown document appropriate for the type:
   - CV: Full academic/professional history, all sections
   - Resume: Concise 1-2 page format, tailored to a role
   - Cover Letter: Professional letter format with date, recipient, body paragraphs, sign-off
5. Format the filename using the rules above:
   - With name: Firstname_Lastname_CV.docx
   - Without name: My_CV.docx, My_Resume.docx, My_Cover_Letter.docx
6. Execute save_to_word — pass your full markdown into text_content and the formatted filename into filename.
7. After the tool succeeds, return TWO things:
   - The FULL document markdown text so the user can read it
   - The exact filename used, in this exact format: Saved as My_CV.docx
8. If no background information was provided by the user at all, ask the user:
   "To generate your document, please share some background information such as your name, experience, skills, and education." 
   Do NOT call save_to_word in this case.""",
    tools=[my_doc_tool, my_search_tool]
)

# High-level orchestrator that routes between the sub-agents based on the user's request.
root_agent = Agent(
    # name="ResearchCoordinator",
    name="Edrick",
    model=my_model,
    instruction="""You are an expert in writting CVs, Resumes & Cover Letters and Doing Research.

Analyze the user's request and follow the correct path:

---
PATH 1: CAREER DOCUMENT GENERATION (CV, Resume, or Cover Letter)
Triggered only when the user asks to create, write, or generate any of: a CV, resume, or cover letter.

Before calling any agent, check if the user has provided any personal background
such as their name, experience, skills, or education.

- If NO background info was provided: Do NOT call any agent. Instead ask the user:
  "To generate your document, I will need some background information.
   Please share details such as your full name, work experience, skills, and education."

- If background info WAS provided: proceed with the steps below:
  1. Call ResearchAgent to gather and structure the user's professional background.
  2. Pass that raw research AND the document type to DocumentStylistAgent.
  3. When DocumentStylistAgent returns, relay to the user:
     - The full document text
     - The confirmation of the filename saved e.g. Saved as Firstname_Lastname_Resume.docx
  DO NOT attempt to format, read, or save the document yourself.
---
PATH 2: GENERAL RESEARCH / SUMMARY
Triggered only when the user asks for research, facts, or a summary on any topic.

1. Call ResearchAgent to gather findings.
2. Pass findings to SummarizerAgent for a concise bulleted summary.
3. Return the final summary to the user.
---
PATH 3: SEARCH GENERAL INFORMATION
Triggered when the user asks for general information, facts, or data that can be answered with a simple search.
1. Use the my_search_tool directly to fetch relevant information. Most especially for current events, facts, or data.
---
""",
    tools=[
        AgentTool(research_agent),
        AgentTool(summarizer_agent),
        AgentTool(cv_stylist_agent),
        (my_search_tool)
        # For anything else, respond helpfully from your own knowledge.
    ],
)

# Wrap your local CLI code in this guard
if __name__ == "__main__":
    # Run the orchestrator locally for rapid terminal debugging
    user_input = input("Enter your request: ")
    runner = InMemoryRunner(agent=root_agent)
    response = runner.run_debug(user_input)
    print(response)

# agent_logic.py — add this function at the bottom

def switch_to_gemini():
    """Rebuild all agents using Gemini as the model."""
    global my_model, research_agent, summarizer_agent, cv_stylist_agent, root_agent

    # my_model = gemini_model
    my_model = models[1]["model"]  # Gemini is the second in the list
    print("[Model Router] Switching all agents to Gemini 2.0 Flash", flush=True)

    research_agent.model = models[1]["model"]  # Gemini
    summarizer_agent.model = models[1]["model"]  # Gemini
    cv_stylist_agent.model = models[1]["model"]  # Gemini
    root_agent.model = models[1]["model"]  # Gemini

def switch_to_groq():
    """Rebuild all agents using Groq as the model."""
    global my_model, research_agent, summarizer_agent, cv_stylist_agent, root_agent

    my_model = models[0]["model"]  # Groq is the first in the list
    print("[Model Router] Switching all agents to Groq llama-3.3-70b", flush=True)

    research_agent.model = models[0]["model"]  # Groq
    summarizer_agent.model = models[0]["model"]  # Groq
    cv_stylist_agent.model = models[0]["model"]  # Groq
    root_agent.model = models[0]["model"]  # Groq